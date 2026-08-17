import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = Document()

    # 页面边距设置 (A4, 2.54cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 样式配置
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Microsoft YaHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal_style.paragraph_format.line_spacing = 1.3
    normal_style.paragraph_format.space_after = Pt(4)

    # 文档大标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("智能机器人管理系统")
    run_title.font.name = 'Microsoft YaHei'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("阶段1（1）多设备并发与远程接入测试报告")
    run_sub.font.name = 'Microsoft YaHei'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # 元数据表格
    table_meta = doc.add_table(rows=3, cols=4)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False

    meta_data = [
        [("测试阶段", True), ("阶段1：模拟接入与通信验证", False), ("测试日期", True), ("2026年8月17日", False)],
        [("测试环境", True), ("本地仿真环境（无公网IP场景）", False), ("并发设备量", True), ("10 台（多型号混合）", False)],
        [("测试人员", True), ("系统开发测试组", False), ("测试结论", True), ("各项通信与调度功能正常", False)]
    ]

    for row_idx, row in enumerate(meta_data):
        for col_idx, (text, is_header) in enumerate(row):
            cell = table_meta.rows[row_idx].cells[col_idx]
            cell.text = text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if is_header:
                set_cell_background(cell, "F1F5F9")
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_heading_1(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return p

    def add_heading_2(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    def add_screenshot_placeholder(box_title, description="请在此处粘贴测试截图（Ctrl+V）："):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run_tag = p.add_run(f"【图表占位：{box_title}】\n")
        run_tag.font.bold = True
        run_tag.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        run_tag.font.size = Pt(10)

        run_desc = p.add_run(f"{description}\n\n\n\n\n（截图粘贴区域）\n")
        run_desc.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        run_desc.font.size = Pt(9.5)

        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_after = Pt(6)

    # 1. 测试背景与环境说明
    add_heading_1("一、 测试背景与环境配置说明")
    p1 = doc.add_paragraph("根据项目阶段1计划，本阶段重点验证多台机器人设备的并发接入、遥测数据周期上报、数据解析入库以及远程控制指令下发链路。")

    add_heading_2("1.1 测试环境与网络架构说明")
    p_net = doc.add_paragraph(
        "由于当前测试环境处于开发与验证阶段，测试机暂未分配独立公网固定 IP 地址，"
        "本次测试采用单机多实体进程架构进行仿真验证：\n"
        "1. 消息中间件（MQTT Broker）运行于本地监听标准端口（1883）；\n"
        "2. 后台服务程序（FastAPI + SQLite）作为独立服务进程接入 Broker；\n"
        "3. 模拟器生成 10 台独立虚拟机器人客户端，分别建立网络会话进行数据推送与指令监听；\n"
        "4. 本地环境验证了协议解析、并发负载与控制时序，与公网 WAN 部署在应用层与传输层协议上保持完全一致。"
    )

    add_heading_2("1.2 跨公网（非局域网）部署方案对照")
    table_env = doc.add_table(rows=4, cols=3)
    table_env.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["配置维度", "本次测试环境（本地模拟）", "生产/现场公网部署方案"]
    for i, h in enumerate(headers):
        cell = table_env.rows[0].cells[i]
        cell.text = h
        set_cell_background(cell, "E2E8F0")
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    env_rows = [
        ["网络寻址方式", "Localhost (127.0.0.1:1883)", "服务器固定公网 IP / 域名 (例: 120.xx.xx.xx:1883)"],
        ["设备接入途径", "同一主机上的独立进程客户端", "异地工厂局域网内的转接电脑发起出网 TCP 连接"],
        ["防火墙与端口", "无需端口映射，直接本机回环", "服务器开放 1883(MQTT)、8000(HTTP/Web) 端口"]
    ]
    for r_idx, row_data in enumerate(env_rows):
        for c_idx, val in enumerate(row_data):
            cell = table_env.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 2. 10台设备测试详情
    add_heading_1("二、 10台设备并发接入与状态呈现测试")
    p_dev = doc.add_paragraph(
        "本次测试模拟接入了 10 台机器人设备，涵盖华数机械臂、立讯 AMR 及巡检机器狗三种型号，设备清单如下："
    )

    table_dev = doc.add_table(rows=11, cols=4)
    table_dev.alignment = WD_TABLE_ALIGNMENT.CENTER
    d_headers = ["设备ID", "设备类型", "上报周期", "测试主要指标"]
    for i, h in enumerate(d_headers):
        cell = table_dev.rows[0].cells[i]
        cell.text = h
        set_cell_background(cell, "E2E8F0")
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dev_list = [
        ["luxshare_amr / amr_001", "搬运 AMR", "3.0 秒", "运行状态、电量、坐标位置、报警码"],
        ["robot_dog / dog_002", "四足巡检狗", "3.0 秒", "巡检步态、电量、姿态角、通信延时"],
        ["huashu_arm / arm_003", "华数机械臂", "3.0 秒", "各轴角度(J1-J6)、末端坐标、程序状态"],
        ["luxshare_amr / amr_004", "搬运 AMR", "3.0 秒", "运行状态、电量、坐标位置、报警码"],
        ["robot_dog / dog_005", "四足巡检狗", "3.0 秒", "巡检步态、电量、姿态角、通信延时"],
        ["huashu_arm / arm_006", "华数机械臂", "3.0 秒", "各轴角度(J1-J6)、末端坐标、程序状态"],
        ["luxshare_amr / amr_007", "搬运 AMR", "3.0 秒", "运行状态、电量、坐标位置、报警码"],
        ["robot_dog / dog_008", "四足巡检狗", "3.0 秒", "巡检步态、电量、姿态角、通信延时"],
        ["huashu_arm / arm_009", "华数机械臂", "3.0 秒", "各轴角度(J1-J6)、末端坐标、程序状态"],
        ["luxshare_amr / amr_010", "搬运 AMR", "3.0 秒", "运行状态、电量、坐标位置、报警码"]
    ]
    for r_idx, row_data in enumerate(dev_list):
        for c_idx, val in enumerate(row_data):
            cell = table_dev.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_heading_2("2.1 测试项 1：Web 大屏 10 台设备概览与卡片渲染")
    doc.add_paragraph("【验证目标】前端页面能正确拉取并渲染 10 台并发设备的卡片，展示在线状态、实时电量及工作模式。")
    add_screenshot_placeholder("Web 大屏 10 台设备卡片概览", "截图指引：截取浏览器 http://127.0.0.1:8000 页面顶部概览统计及 10 张设备卡片完整展示。")

    add_heading_2("2.2 测试项 2：MQTT 消息中间件与后台服务连接状态")
    doc.add_paragraph("【验证目标】页面右上角 MQTT 服务显示为「已连接」，后台控制台正常输出主题订阅及并发入库日志。")
    add_screenshot_placeholder("MQTT 连接状态与后台运行日志", "截图指引：截取网页右上角「MQTT 服务: 已连接」绿灯标识，或终端后台收发报文日志。")

    # 3. 数据解析与入库测试
    add_heading_1("三、 实时遥测数据与传感器数据解析测试")
    add_heading_2("3.1 测试项 3：华数机械臂关节与末端数据上报")
    doc.add_paragraph("【验证目标】选中华数机械臂设备时，下方的遥测面板能解析出 J1~J6 六轴角度数据及 X/Y/Z 笛卡尔空间坐标。")
    add_screenshot_placeholder("华数机械臂多轴数据解析界面", "截图指引：在页面点击任一华数机械臂卡片（如 arm_003），截取下方「遥测参数」中显示的六轴角度及末端坐标。")

    add_heading_2("3.2 测试项 4：原始 MQTT 报文入库与 JSON 格式解析")
    doc.add_paragraph("【验证目标】后台成功接收规范格式的 JSON 报文并入库，前端可查看到结构化的原始上报数据。")
    add_screenshot_placeholder("原始 MQTT 报文 JSON 视图", "截图指引：截取设备详情面板中的「原始 MQTT 报文 (JSON)」查看器内容。")

    # 4. 远程控制测试
    add_heading_1("四、 远程控制指令下发与响应测试")
    add_heading_2("4.1 测试项 5：任务控制指令发布 (cmd/# 主题)")
    doc.add_paragraph("【验证目标】在 Web 控制面板选择控制动作（如启动、急停、复位）并点击下发，指令通过 MQTT 路由至对应设备主题并被消费。")
    add_screenshot_placeholder("Web 远程指令下发操作与反馈", "截图指引：在右下角「远程控制与任务调度」选择指令并点击确认下发，截取提示成功的 Toast 弹窗或终端反馈日志。")

    # 5. 心跳与离线
    add_heading_1("五、 心跳超时与设备离线判定测试")
    add_heading_2("5.1 测试项 6：自动离线检测机制")
    doc.add_paragraph("【验证目标】当模拟设备停止上报超过超时阈值（15秒）后，系统自动将设备状态置为「离线」，页面指示灯变为灰色。")
    add_screenshot_placeholder("设备停止上报后自动离线状态", "截图指引：终止模拟器进程约 15 秒后，截取网页上设备状态变为「离线」的展示图。")

    # 6. 测试结论与后续工作
    add_heading_1("六、 测试结论与阶段后续工作")
    p_concl = doc.add_paragraph(
        "1. 本次测试在本地仿真环境下顺利完成了 10 台跨类型机器人的并发接入与双向通信验证；\n"
        "2. MQTT 发布/订阅机制、JSON 数据解析与 SQLite 高频写入链路工作平稳，未出现阻塞或丢包异常；\n"
        "3. 下一步将根据《华数SDK对接要求》开展转接程序（Adapter Bridge）的编写与现场实体机器人对接联调。"
    )

    desktop_path = r"C:\Users\32755\Desktop\智能机器人管理系统_多设备并发与远程接入测试报告.docx"
    doc.save(desktop_path)
    print(f"Document successfully created at {desktop_path}")

if __name__ == "__main__":
    create_document()
